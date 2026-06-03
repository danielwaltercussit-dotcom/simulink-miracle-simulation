from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
RAW_JSON = ROOT / "build" / "reports" / "github_simulink_skill_research_raw.json"
OUT = ROOT / "Simulink建模Skills调研与推荐清单_2026-05-29.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_fonts(run, east_asia: str = "宋体", ascii_font: str = "Calibri") -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_fonts(r)
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_fonts(run, "黑体", "Calibri")


def repo(raw: list[dict], name: str) -> dict:
    for item in raw:
        if item["full_name"].lower() == name.lower():
            return item
    return {
        "full_name": name,
        "stars": "",
        "forks": "",
        "updated": "",
        "description": "",
        "url": f"https://github.com/{name}",
    }


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "D9EAF7")
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        set_fonts(run, "黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            run = p.add_run(str(value))
            set_fonts(run)


def main() -> None:
    raw = json.loads(RAW_JSON.read_text(encoding="utf-8-sig"))

    key_names = [
        "matlab/simulink-agentic-toolkit",
        "matlab/agent-skills-playground",
        "simulink/skills",
        "npuzsy/simulink-power-electronics",
        "efantnu/pwrsys-matlab",
        "mathworks/Simscape_Electrical_Support_Library",
        "McSCert/Simulink-Utility",
        "sohumsuthar/simulink-mcp",
        "McSCert/Auto-Layout",
        "huanhyougo/matlab-power-electronics-skill",
        "K-Dense-AI/scientific-agent-skills",
    ]
    repos = [repo(raw, name) for name in key_names]

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GitHub Simulink 建模 Skills 调研与当前项目推荐清单")
    run.bold = True
    run.font.size = Pt(18)
    set_fonts(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("调研日期：2026-05-29；项目：simulink_agent_v1；目标：IEEE39/DFIG 电力电子化电力系统建模")
    set_fonts(run)

    add_heading(doc, "一、结论摘要", 1)
    conclusions = [
        "最高优先级仍然是 MathWorks 官方 Simulink Agentic Toolkit：它同时提供 MCP 工具和 Model-Based Design skills，适合作为本项目建模、编辑、仿真、测试的基础能力。",
        "simulink/skills 是官方工具包之外的专门工作流补充，适合模型交互、命令行调试、初始化与仿真性能分析。",
        "电力电子方向新增值得关注的是 npuzsy/simulink-power-electronics：星标不算最高，但它是直接面向 Codex、Simulink、Simscape Electrical 和电力电子布局/验证的 skill，和本项目 DFIG/并网变换器场景贴合度高。",
        "huanhyougo/matlab-power-electronics-skill 覆盖模型很多，但提交历史很短、星标较低，建议先当作案例库和术语/模型清单参考，暂不直接作为主工作流依据。",
        "McSCert/Auto-Layout 和 Simulink-Utility 对排版有用，但对电力系统物理网络不能全局自动排版；只建议用于普通 Simulink 控制/测量子系统或派生副本。",
    ]
    for item in conclusions:
        add_paragraph(doc, item, "List Bullet")

    add_heading(doc, "二、GitHub 候选库排名与用途", 1)
    rows = []
    for item in repos:
        rows.append(
            [
                item["full_name"],
                str(item.get("stars", "")),
                str(item.get("forks", "")),
                str(item.get("updated", ""))[:10],
                item.get("description") or "",
                item.get("url") or "",
            ]
        )
    add_table(doc, ["仓库", "Stars", "Forks", "更新日期", "说明", "链接"], rows)

    add_paragraph(
        doc,
        "注：Stars/Forks 来自 GitHub API 在线检索结果；GitHub 网页缓存显示可能与 API 实时值有少量差异。",
    )

    add_heading(doc, "三、当前项目建议采用的 Skills 组合", 1)
    recommended = [
        [
            "building-simulink-models",
            "主建模/编辑",
            "新增、替换、连接 blocks；修改参数；编辑 Simscape/Stateflow/System Composer 结构。",
            "每次先读模型，再按子系统局部编辑；避免手写脆弱 block path。",
        ],
        [
            "simulating-simulink-models",
            "仿真与数据分析",
            "用 SimulationInput/SimulationOutput 做短仿真、参数扫掠、日志读取。",
            "IEEE39/DFIG 场景至少保留 update + 0.005s smoke。",
        ],
        [
            "testing-simulink-models",
            "回归测试",
            "为关键子系统或行为写可复用 pass/fail 测试。",
            "当模型稳定后，把 smoke test 固化成测试资产。",
        ],
        [
            "simulink-auto-layout-github",
            "排版与连线",
            "根据 GitHub Auto-Layout、Simulink-Utility 和内置 routeLine/arrangeSystem 做布局辅助。",
            "电气物理网络用确定性坐标；自动排版只给普通控制/测量子系统。",
        ],
        [
            "power-electronics-component-libraries",
            "电力电子组件选择",
            "在 DFIG、VSC、MMC、PLL、PI、滤波器、机器模型等来源间做兼容性选择。",
            "先保留 NEBUS benchmark，再叠加 converter scenario overlay。",
        ],
        [
            "simulink-interactions",
            "当前打开模型交互",
            "定位当前模型、子系统、选中 block，并做小范围交互式修改。",
            "适合 MATLAB 已打开目标模型时使用。",
        ],
        [
            "simulink-debug-commandline",
            "仿真故障定位",
            "命令行检查特定 block、特定时刻和仿真内部行为。",
            "用于 update/sim 报错或波形异常的精确诊断。",
        ],
        [
            "simulink-profile-initialization / profiler / solver-profiler",
            "性能诊断",
            "分析初始化、仿真、solver 性能瓶颈。",
            "模型规模继续膨胀或编译过慢时启用。",
        ],
        [
            "code-simplifier",
            "脚本清理",
            "清理生成脚本、验证脚本和报告生成脚本。",
            "保留功能，减少硬编码和重复路径。",
        ],
        [
            "find-skill / skill-creator / document-skills",
            "技能检索、沉淀与文档",
            "发现外部 skills、把项目经验固化成 skill、生成 Word/PDF/PPT 文档。",
            "本报告即由 document-skills 工作流生成。",
        ],
    ]
    add_table(doc, ["Skill", "定位", "用途", "本项目用法"], recommended)

    add_heading(doc, "四、对 IEEE39/DFIG 建模的执行规则", 1)
    rules = [
        "NEBUS39V2.slx 始终作为 benchmark oracle；不要直接改标准模型，只派生生成模型。",
        "替换 SG 为 DFIG 时必须记录替换关系、bus、machine row、源 block 和 scenario_role。",
        "布局优先沿用标准模型坐标；局部遮挡只做确定性偏移；物理电气连接不使用 Goto/From 替代。",
        "电力电子组件先查兼容性，再引入：R2024b 当前优先用已有 DFIG 模板和本地可运行库，新版 Simscape 组件先作为参考。",
        "每次建模完成必须出报告：DFIG 数量、剩余 SG 数量、根层 overlap、SimulationCommand update、短时 sim smoke。",
        "GitHub 外部库只装到项目本地 external/github 或 .agents/skills，不写全局 MATLAB/Codex 配置，除非用户明确要求。",
    ]
    for item in rules:
        add_paragraph(doc, item, "List Number")

    add_heading(doc, "五、是否建议新增安装", 1)
    add_paragraph(
        doc,
        "建议先不替换当前主工作流。当前项目本地已有 MathWorks MBD skills、simulink-skills、auto-layout、power-electronics-component-libraries 等核心能力。若下一步要专门做并网逆变器、MMC、STATCOM、SVPWM/VSG 或电力电子波形诊断，可优先试装 npuzsy/simulink-power-electronics 到项目本地 .agents/skills，并用其校验脚本检查结构。",
    )
    add_paragraph(
        doc,
        "huanhyougo/matlab-power-electronics-skill 可以作为案例和模型清单参考；因为提交历史短且星标低，不建议直接把它作为自动建模标准。",
    )

    add_heading(doc, "六、主要来源", 1)
    sources = [
        "https://github.com/matlab/simulink-agentic-toolkit",
        "https://github.com/simulink/skills",
        "https://github.com/matlab/agent-skills-playground",
        "https://github.com/npuzsy/simulink-power-electronics",
        "https://github.com/huanhyougo/matlab-power-electronics-skill",
        "https://github.com/McSCert/Auto-Layout",
        "https://github.com/McSCert/Simulink-Utility",
        "https://github.com/mathworks/Simscape_Electrical_Support_Library",
        "https://github.com/efantnu/pwrsys-matlab",
        "https://github.com/sohumsuthar/simulink-mcp",
    ]
    for source in sources:
        add_paragraph(doc, source, "List Bullet")

    doc.add_page_break()
    add_heading(doc, "附录：本次检索说明", 1)
    add_paragraph(
        doc,
        "检索方式：GitHub 搜索关键词包括 simulink skills、matlab simulink agent skills、simulink agentic toolkit、simulink auto layout、matlab simulink power systems converter interfaced、matlab mcp simulink、simulink model based design skills；同时打开关键仓库 README 复核用途。",
    )
    add_paragraph(
        doc,
        f"原始检索数据保存在：{RAW_JSON.relative_to(ROOT)}。文档生成日期：{date.today().isoformat()}。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
